from fastapi import FastAPI, HTTPException, Depends, status, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from datetime import datetime, timedelta
from passlib.context import CryptContext
from jose import JWTError, jwt
import os
import uuid
import logging
import pandas as pd
import io
from docx import Document
from docx.shared import Inches
import re
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# MongoDB setup
MONGO_URL = os.getenv("MONGO_URL", "mongodb+srv://royalcert_user:Ccpp1144..@royalcert-cluster.l1hqqn.mongodb.net/")
DATABASE_NAME = "royalcert_db"

client = AsyncIOMotorClient(MONGO_URL)
db = client[DATABASE_NAME]

# JWT Setup
SECRET_KEY = os.getenv("SECRET_KEY", "your-super-secret-key-change-this-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 1440  # 24 hours

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Security
security = HTTPBearer()

app = FastAPI(title="RoyalCert Periyodik Muayene Sistemi", version="1.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify exact origins
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# ===================== MODELS =====================

class UserRole(str):
    ADMIN = "admin"
    PLANLAMA_UZMANI = "planlama_uzmani" 
    DENETCI = "denetci"
    TEKNIK_YONETICI = "teknik_yonetici"

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    email: str
    full_name: str
    role: str
    is_active: bool = True
    created_at: datetime = Field(default_factory=datetime.utcnow)

class UserCreate(BaseModel):
    username: str
    email: str
    full_name: str
    password: str
    role: str

class UserLogin(BaseModel):
    username: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str
    user: User

class Customer(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    contact_person: str
    phone: str
    email: str
    address: str
    equipments: List[Dict[str, Any]] = []
    created_at: datetime = Field(default_factory=datetime.utcnow)

class CustomerCreate(BaseModel):
    company_name: str
    contact_person: str
    phone: str
    email: str
    address: str
    equipments: List[Dict[str, Any]] = []

class ControlItem(BaseModel):
    id: int
    text: str
    category: str
    input_type: str = "dropdown"  # dropdown, text, number
    has_comment: bool = True
    required: bool = True

class EquipmentCategory(BaseModel):
    code: str  # A, B, C, D, E, F, G, H
    name: str
    items: List[ControlItem]

class EquipmentTemplate(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: Optional[str] = None
    equipment_type: str
    template_type: Optional[str] = None  # FORM or REPORT
    description: str
    categories: List[EquipmentCategory]
    created_by: str
    is_active: bool = True
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class EquipmentTemplateCreate(BaseModel):
    name: Optional[str] = None
    equipment_type: str
    template_type: Optional[str] = None  # FORM or REPORT
    description: str
    categories: List[EquipmentCategory]

class Inspection(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    customer_id: str
    equipment_info: Dict[str, Any]
    inspector_id: str
    planned_date: datetime
    status: str = "beklemede"  # beklemede, devam_ediyor, rapor_yazildi, onaylandi, reddedildi
    report_data: Optional[Dict[str, Any]] = None  # Denetim rapor verileri
    approval_notes: Optional[str] = None  # Teknik yönetici notları
    approved_by: Optional[str] = None  # Onaylayan teknik yöneticinin ID'si
    approved_at: Optional[datetime] = None  # Onaylanma tarihi
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class InspectionCreate(BaseModel):
    customer_id: str
    equipment_info: Dict[str, Any]
    inspector_id: str
    planned_date: datetime

class InspectionCreate(BaseModel):
    customer_id: str
    equipment_info: Dict[str, Any]
    inspector_id: str
    planned_date: datetime

class InspectionUpdate(BaseModel):
    status: Optional[str] = None
    report_data: Optional[Dict[str, Any]] = None

class InspectionApproval(BaseModel):
    action: str  # "approve" or "reject"
    notes: Optional[str] = None

class BulkImportResult(BaseModel):
    total_rows: int
    successful_imports: int
    failed_imports: int
    errors: List[Dict[str, Any]] = []
    warnings: List[str] = []

class BulkImportItem(BaseModel):
    muayene_alani: Optional[str] = None
    muayene_alt_alani: Optional[str] = None
    muayene_turu: Optional[str] = None
    referans: Optional[str] = None
    muayene_tarihi: Optional[str] = None
    zorunlu_alan: Optional[str] = None
    musteri_adi: str
    musteri_adresi: str
    denetci_adi: Optional[str] = None
    denetci_lokasyonu: Optional[str] = None
    rapor_onay_tarihi: Optional[str] = None
    rapor_onaylayan: Optional[str] = None

# ===================== AUTH FUNCTIONS =====================

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = await db.users.find_one({"username": username})
    if user is None:
        raise credentials_exception
    
    return User(**user)

def require_role(required_role: str):
    def role_checker(current_user: User = Depends(get_current_user)):
        if current_user.role != required_role and current_user.role != UserRole.ADMIN:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not enough permissions"
            )
        return current_user
    return role_checker

# ===================== API ENDPOINTS =====================

@app.post("/api/auth/register", response_model=User)
async def register(user_create: UserCreate, current_user: User = Depends(require_role(UserRole.ADMIN))):
    # Check if username already exists
    existing_user = await db.users.find_one({"username": user_create.username})
    if existing_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    # Hash password
    hashed_password = get_password_hash(user_create.password)
    
    # Create user
    user_dict = user_create.dict()
    user_dict["password"] = hashed_password
    user_dict["id"] = str(uuid.uuid4())
    user_dict["created_at"] = datetime.utcnow()
    
    await db.users.insert_one(user_dict)
    
    # Remove password from response
    del user_dict["password"]
    return User(**user_dict)

@app.post("/api/auth/login", response_model=Token)
async def login(user_login: UserLogin):
    user = await db.users.find_one({"username": user_login.username})
    
    if not user or not verify_password(user_login.password, user["password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    if not user.get("is_active", True):
        raise HTTPException(status_code=400, detail="Inactive user")
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["username"]}, expires_delta=access_token_expires
    )
    
    # Remove password from user data
    del user["password"]
    
    return Token(
        access_token=access_token,
        token_type="bearer",
        user=User(**user)
    )

@app.get("/api/auth/me", response_model=User)
async def get_current_user_info(current_user: User = Depends(get_current_user)):
    return current_user

# ===================== USER MANAGEMENT =====================

@app.get("/api/users", response_model=List[User])
async def get_users(current_user: User = Depends(require_role(UserRole.ADMIN))):
    users = await db.users.find({}, {"password": 0}).to_list(1000)
    return [User(**user) for user in users]

@app.put("/api/users/{user_id}")
async def update_user(user_id: str, user_data: dict, current_user: User = Depends(require_role(UserRole.ADMIN))):
    # Build update data
    update_data = {}
    
    if "full_name" in user_data:
        update_data["full_name"] = user_data["full_name"]
    if "email" in user_data:
        update_data["email"] = user_data["email"] 
    if "username" in user_data:
        update_data["username"] = user_data["username"]
    if "role" in user_data:
        update_data["role"] = user_data["role"]
    if "is_active" in user_data:
        update_data["is_active"] = user_data["is_active"]
    
    # Update password if provided
    if "password" in user_data and user_data["password"]:
        hashed_password = pwd_context.hash(user_data["password"])
        update_data["password_hash"] = hashed_password
    
    update_data["updated_at"] = datetime.utcnow()
    
    result = await db.users.update_one(
        {"id": user_id}, 
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Return updated user
    updated_user = await db.users.find_one({"id": user_id})
    return User(**updated_user)

@app.put("/api/users/{user_id}/password")
async def change_user_password(user_id: str, password_data: dict, current_user: User = Depends(require_role(UserRole.ADMIN))):
    if "new_password" not in password_data:
        raise HTTPException(status_code=400, detail="new_password field is required")
    
    new_password = password_data["new_password"]
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters long")
    
    hashed_password = pwd_context.hash(new_password)
    
    result = await db.users.update_one(
        {"id": user_id},
        {"$set": {
            "password_hash": hashed_password,
            "updated_at": datetime.utcnow()
        }}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "Password changed successfully"}

@app.delete("/api/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(require_role(UserRole.ADMIN))):
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    return {"message": "User deleted successfully"}

# ===================== CUSTOMER MANAGEMENT =====================

@app.post("/api/customers", response_model=Customer)
async def create_customer(customer: CustomerCreate, current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))):
    customer_dict = customer.dict()
    customer_dict["id"] = str(uuid.uuid4())
    customer_dict["created_at"] = datetime.utcnow()
    
    await db.customers.insert_one(customer_dict)
    return Customer(**customer_dict)

@app.get("/api/customers", response_model=List[Customer])
async def get_customers(current_user: User = Depends(get_current_user)):
    customers = await db.customers.find().to_list(1000)
    return [Customer(**customer) for customer in customers]

@app.get("/api/customers/{customer_id}", response_model=Customer)
async def get_customer(customer_id: str, current_user: User = Depends(get_current_user)):
    customer = await db.customers.find_one({"id": customer_id})
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    return Customer(**customer)

@app.put("/api/customers/{customer_id}", response_model=Customer)
async def update_customer(customer_id: str, customer_update: CustomerCreate, current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))):
    result = await db.customers.update_one(
        {"id": customer_id},
        {"$set": customer_update.dict()}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    updated_customer = await db.customers.find_one({"id": customer_id})
    return Customer(**updated_customer)

@app.delete("/api/customers/{customer_id}")
async def delete_customer(customer_id: str, current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))):
    result = await db.customers.delete_one({"id": customer_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Customer not found")
    return {"message": "Customer deleted successfully"}

# ===================== BULK IMPORT =====================

@app.post("/api/customers/bulk-import", response_model=BulkImportResult)
async def bulk_import_customers(
    file: UploadFile = File(...),
    current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))
):
    """
    Excel dosyasından müşteri toplu yükleme
    Excel sütunları: A: Muayene Alanı, B: Muayene Alt Alanı, C: Muayene Türü, 
    D: Referans, E: Muayene Tarihi, F: Zorunlu/Gönüllü Alan, G: Müşteri Adı, 
    H: Müşteri Adresi, I: Denetçi Adı, J: Denetçi Lokasyonu, 
    K: Rapor Onay Tarihi, L: Raporu Onaylayan Teknik Yönetici
    """
    
    # File type validation
    if not file.filename.endswith(('.xlsx', '.xls', '.csv')):
        raise HTTPException(status_code=400, detail="Sadece Excel (.xlsx, .xls) veya CSV dosyaları kabul edilir")
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Parse Excel/CSV
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(file_content))
        else:
            df = pd.read_excel(io.BytesIO(file_content), engine='openpyxl')
        
        # Define expected columns
        expected_columns = [
            'Muayene Alanı',
            'Muayene Alt Alanı', 
            'Muayene Türü',
            'Referans',
            'Muayene Tarihi',
            'Zorunlu Alan ya da Gönüllü Alan',
            'Müşteri Adı',
            'Müşteri Adresi',
            'Denetçi Adı',
            'Denetçinin Lokasyonu',
            'Rapor Onay Tarihi',
            'Raporu Onaylayan Teknik Yönetici'
        ]
        
        # Check if we have the expected columns, if not use first 12 columns
        if len(df.columns) >= 12:
            df.columns = expected_columns[:len(df.columns)]
        else:
            raise HTTPException(status_code=400, detail="Excel dosyasında en az 12 sütun olmalı")
        
        # Initialize counters
        total_rows = len(df)
        successful_imports = 0
        failed_imports = 0
        errors = []
        warnings = []
        
        # Process each row
        for index, row in df.iterrows():
            try:
                # Skip empty rows or rows without customer name
                customer_name = str(row.get('Müşteri Adı', '')).strip()
                customer_address = str(row.get('Müşteri Adresi', '')).strip()
                
                if not customer_name or customer_name == 'nan' or customer_name == '-':
                    warnings.append(f"Satır {index + 2}: Müşteri adı boş, atlanıyor")
                    continue
                
                if not customer_address or customer_address == 'nan' or customer_address == '-':
                    warnings.append(f"Satır {index + 2}: Müşteri adresi boş, atlanıyor") 
                    continue
                
                # Clean and prepare data
                def clean_value(val):
                    if pd.isna(val) or str(val).strip() in ['', 'nan', '-']:
                        return None
                    return str(val).strip()
                
                # Check if customer already exists
                existing_customer = await db.customers.find_one({"company_name": customer_name})
                
                equipment_info = {}
                muayene_alani = clean_value(row.get('Muayene Alanı'))
                muayene_alt_alani = clean_value(row.get('Muayene Alt Alanı'))
                
                if muayene_alani or muayene_alt_alani:
                    equipment_info = {
                        "muayene_alani": muayene_alani,
                        "muayene_alt_alani": muayene_alt_alani,
                        "muayene_turu": clean_value(row.get('Muayene Türü')),
                        "referans": clean_value(row.get('Referans')),
                        "muayene_tarihi": clean_value(row.get('Muayene Tarihi')),
                        "zorunlu_alan": clean_value(row.get('Zorunlu Alan ya da Gönüllü Alan')),
                        "denetci_adi": clean_value(row.get('Denetçi Adı')),
                        "denetci_lokasyonu": clean_value(row.get('Denetçinin Lokasyonu')),
                        "rapor_onay_tarihi": clean_value(row.get('Rapor Onay Tarihi')),
                        "rapor_onaylayan": clean_value(row.get('Raporu Onaylayan Teknik Yönetici'))
                    }
                
                if existing_customer:
                    # Update existing customer with new equipment info if provided
                    if equipment_info:
                        existing_equipments = existing_customer.get('equipments', [])
                        
                        # Check if similar equipment already exists
                        equipment_exists = False
                        for eq in existing_equipments:
                            if (eq.get('muayene_alani') == equipment_info.get('muayene_alani') and 
                                eq.get('muayene_alt_alani') == equipment_info.get('muayene_alt_alani')):
                                equipment_exists = True
                                break
                        
                        if not equipment_exists:
                            existing_equipments.append(equipment_info)
                            await db.customers.update_one(
                                {"id": existing_customer["id"]},
                                {"$set": {"equipments": existing_equipments}}
                            )
                            warnings.append(f"Satır {index + 2}: Müşteri mevcut, yeni ekipman bilgisi eklendi")
                        else:
                            warnings.append(f"Satır {index + 2}: Müşteri ve ekipman bilgisi zaten mevcut")
                else:
                    # Create new customer
                    customer_data = {
                        "id": str(uuid.uuid4()),
                        "company_name": customer_name,
                        "contact_person": customer_name,  # Default to company name
                        "phone": "",  # Will be empty, can be updated later
                        "email": "",  # Will be empty, can be updated later
                        "address": customer_address,
                        "equipments": [equipment_info] if equipment_info else [],
                        "created_at": datetime.utcnow(),
                        "import_source": "bulk_import",
                        "imported_by": current_user.id,
                        "imported_at": datetime.utcnow()
                    }
                    
                    await db.customers.insert_one(customer_data)
                
                successful_imports += 1
                    
            except Exception as e:
                failed_imports += 1
                errors.append({
                    "row": index + 2,
                    "error": str(e),
                    "data": dict(row) if hasattr(row, 'to_dict') else str(row)
                })
        
        result = BulkImportResult(
            total_rows=total_rows,
            successful_imports=successful_imports,
            failed_imports=failed_imports,
            errors=errors,
            warnings=warnings
        )
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Dosya işleme hatası: {str(e)}")

@app.get("/api/customers/bulk-import/template")
async def download_bulk_import_template(current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))):
    """Bulk import için Excel template'ini indirme endpoint'i"""
    
    # Template headers
    headers = [
        'Muayene Alanı',
        'Muayene Alt Alanı', 
        'Muayene Türü',
        'Referans',
        'Muayene Tarihi',
        'Zorunlu Alan ya da Gönüllü Alan',
        'Müşteri Adı',
        'Müşteri Adresi',
        'Denetçi Adı',
        'Denetçinin Lokasyonu',
        'Rapor Onay Tarihi',
        'Raporu Onaylayan Teknik Yönetici'
    ]
    
    # Sample data
    sample_data = [
        [
            'Kaldırma ve İndirme Ekipmanları',
            'CARASKAL',
            'PERİYODİK',
            'TSE EN 280',
            '2025-01-15',
            'Zorunlu Alan',
            'ABC İnşaat Ltd. Şti.',
            'İstanbul, Türkiye',
            'Mehmet Yılmaz',
            'İstanbul',
            '2025-01-20',
            'Ali Koç'
        ],
        [
            'İş Güvenliği Ekipmanları',
            'İSKELE',
            'İLK MONTAJ',
            'TS 498',
            '2025-02-10',
            'Gönüllü Alan',
            'XYZ Yapı A.Ş.',
            'Ankara, Türkiye',
            'Ayşe Demir',
            'Ankara',
            '',
            ''
        ]
    ]
    
    # Create DataFrame
    df = pd.DataFrame(sample_data, columns=headers)
    
    # Create Excel file in memory
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Müşteri Listesi')
        
        # Get workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Müşteri Listesi']
        
        # Adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    
    return {
        "message": "Template hazırlandı",
        "filename": "musteri_listesi_template.xlsx",
        "content": output.getvalue().hex()  # Send as hex string
    }

# ===================== EQUIPMENT TEMPLATES =====================

# ===================== CARASKAL TEMPLATE DATA =====================

def get_caraskal_template():
    """Caraskal muayene formu template'i"""
    caraskal_categories = [
        {
            "code": "A",
            "name": "KUMANDA SİSTEMLERİ VE/VEYA KABİN",
            "items": [
                {"id": 1, "text": "Kumandalar ve işaretlemeleri (araba ve köprü hareketleri, kaldırma ve indirme ve kablosuz kumanda kullanımı)", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 2, "text": "El zinciri ve çalışması", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 3, "text": "El kumandası kablo uzunluğunun operatörün tehlike bölgesinden uzaklaşmasını sağlaması ve gergi teli", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 4, "text": "Farklı kumanda konsollarının aynı anda çalışmasının engellenmesi", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 5, "text": "Kabin ve kabin bağlantıları", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 6, "text": "Acil durdurma butonu", "category": "A", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "B",
            "name": "SINIRLAYICILAR",
            "items": [
                {"id": 7, "text": "Nominal kapasite sınırlayıcısı (aşırı yük) sistemi ve ikazı", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 8, "text": "Alt sınır güvenlik sınırlayıcısı", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 9, "text": "Üst sınır güvenlik sınırlayıcısı", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 10, "text": "Araba hareket sınırlayıcılar (tamponlar) ve yavaşlatıcılar", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 11, "text": "Köprü hareket sınırlayıcılar (tamponlar) ve yavaşlatıcılar", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 12, "text": "Kabin hareket sınırlayıcılar (tamponlar) ve yavaşlatıcılar", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 13, "text": "Aynı kılavuz ray üzerinde çalışan caraskalların çarpışmasının engellenmesi", "category": "B", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "C",
            "name": "FREN SİSTEMİ",
            "items": [
                {"id": 14, "text": "Kaldırma ve indirme frenleri", "category": "C", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 15, "text": "Araba hareketinin durması", "category": "C", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 16, "text": "Köprü hareketinin durması", "category": "C", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "D",
            "name": "YÜK ZİNCİRİ, KANCA BLOĞU VE ÇARKLAR",
            "items": [
                {"id": 17, "text": "Yük zincir; Bakla kalıcı uzama", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 18, "text": "Yük zinciri yapısı, deformasyon ve toplanması", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 19, "text": "Kanca gövde kesiti, ağız açıklığı, serbest dönüş ve deformasyon", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 20, "text": "Kanca/kanca bloğu için; yükün istemsiz hareketinin engellenmesi", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 21, "text": "Zincirin çark yuvalarından çıkması ve zincir çark arası el sıkışma tehlikesi", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 22, "text": "Çarklar", "category": "D", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "E",
            "name": "YAPISAL KONTROL",
            "items": [
                {"id": 23, "text": "Yürüme rayları (araba ve köprü) yapısı ve deformasyon", "category": "E", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 24, "text": "Köprü kiriş yapısı ve deformasyon", "category": "E", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 25, "text": "Köprü taşıyıcı kiriş ve kolon yapısı ve deformasyon", "category": "E", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "F",
            "name": "ERİŞİM",
            "items": [
                {"id": 26, "text": "Hareketli parçalar arasında ezilme ve geçit genişliği", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 27, "text": "Çalışma alanlarına ve/veya kabine çıkış merdiveni", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 28, "text": "Dinlenme platformu", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 29, "text": "Çemberli koruyucu", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 30, "text": "Yürüme yolları", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 31, "text": "Erişimde üç nokta desteği", "category": "F", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "G",
            "name": "DİĞER KONTROLLER",
            "items": [
                {"id": 32, "text": "Üretici ve kapasite etiketi", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 33, "text": "(varsa) Aparat kapasite etiketi", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 34, "text": "Çalışma alanı içerisinde uyarı etiket ve işaretleri", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 35, "text": "Sesli ve/veya görsel ikaz", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 36, "text": "Kılavuz raylar arasına yabancı madde girmesinin engellenmesi", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 37, "text": "Redüktör yağlanması ve sızıntı", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 38, "text": "Kilitlenebilir ana şalter", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 39, "text": "Kaçak akıma karşı koruma", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 40, "text": "Kablo arabası ve bükülebilen kablolar", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 41, "text": "Bakım onarım kayıtları, işletme talimatı ve el kitabı", "category": "G", "input_type": "dropdown", "has_comment": True, "required": True}
            ]
        },
        {
            "code": "H",
            "name": "TEST, DENEY VE MUAYENE",
            "items": [
                {"id": 42, "text": "Caraskalın yüksüz durumda testi (fonksiyon testi)", "category": "H", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 43, "text": "Caraskal kaldırma kapasitesinin 1,25 / 1,5 katında test yükü ile gerçekleştirilen test (statik test)", "category": "H", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 44, "text": "Statik testin gerçekleştirildiği test yükü (kg)", "category": "H", "input_type": "text", "has_comment": False, "required": True},
                {"id": 45, "text": "Caraskal kaldırma kapasitesinin 1,1 katında test yükü ile gerçekleştirilen test (dinamik test)", "category": "H", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 46, "text": "Dinamik testin gerçekleştirildiği test yükü (kg)", "category": "H", "input_type": "text", "has_comment": False, "required": True},
                {"id": 47, "text": "Caraskalın işletme/kaldırma kapasitesinde test yükü ile gerçekleştirilen test (yük testi)", "category": "H", "input_type": "dropdown", "has_comment": True, "required": True},
                {"id": 48, "text": "Yük testinin gerçekleştirildiği test yükü (kg.)", "category": "H", "input_type": "text", "has_comment": False, "required": True}
            ]
        }
    ]
    
    return {
        "equipment_type": "CARASKAL",
        "description": "Caraskal Muayene Formu ve Raporu",
        "categories": caraskal_categories
    }

# ===================== TEMPLATE MANAGEMENT API =====================

@app.post("/api/equipment-templates", response_model=EquipmentTemplate)
async def create_equipment_template(
    template: EquipmentTemplateCreate, 
    current_user: User = Depends(require_role(UserRole.ADMIN))
):
    # Check if template already exists
    existing = await db.equipment_templates.find_one({"equipment_type": template.equipment_type})
    if existing:
        raise HTTPException(status_code=400, detail="Equipment template already exists")
    
    template_dict = template.dict()
    template_dict["id"] = str(uuid.uuid4())
    template_dict["created_by"] = current_user.id
    template_dict["is_active"] = True
    template_dict["created_at"] = datetime.utcnow()
    template_dict["updated_at"] = datetime.utcnow()
    
    await db.equipment_templates.insert_one(template_dict)
    return EquipmentTemplate(**template_dict)

@app.get("/api/equipment-templates", response_model=List[EquipmentTemplate])
async def get_equipment_templates(current_user: User = Depends(get_current_user)):
    templates = await db.equipment_templates.find({"is_active": True}).to_list(1000)
    return [EquipmentTemplate(**template) for template in templates]

@app.get("/api/equipment-templates/{template_id}", response_model=EquipmentTemplate)
async def get_equipment_template(template_id: str, current_user: User = Depends(get_current_user)):
    template = await db.equipment_templates.find_one({"id": template_id, "is_active": True})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return EquipmentTemplate(**template)

@app.put("/api/equipment-templates/{template_id}", response_model=EquipmentTemplate)
async def update_equipment_template(
    template_id: str,
    template_update: EquipmentTemplateCreate,
    current_user: User = Depends(require_role(UserRole.ADMIN))
):
    existing = await db.equipment_templates.find_one({"id": template_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Template not found")
    
    update_data = template_update.dict()
    update_data["updated_at"] = datetime.utcnow()
    
    await db.equipment_templates.update_one({"id": template_id}, {"$set": update_data})
    
    updated_template = await db.equipment_templates.find_one({"id": template_id})
    return EquipmentTemplate(**updated_template)

@app.delete("/api/equipment-templates/{template_id}")
async def delete_equipment_template(
    template_id: str,
    current_user: User = Depends(require_role(UserRole.ADMIN))
):
    result = await db.equipment_templates.update_one(
        {"id": template_id},
        {"$set": {"is_active": False, "updated_at": datetime.utcnow()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"message": "Template deleted successfully"}

class InspectionFormResult(BaseModel):
    item_id: int
    category: str
    value: str  # "U", "UD", "UY"
    comment: Optional[str] = None

class InspectionFormData(BaseModel):
    general_info: Dict[str, Any]  # Genel bilgiler (müşteri, tarih vs.)
    equipment_info: Dict[str, Any]  # Ekipman bilgileri
    measurement_tools: List[Dict[str, Any]] = []  # Ölçüm aletleri
    form_results: List[InspectionFormResult]  # Kontrol sonuçları
    defects: Optional[str] = None  # Kusur açıklamaları
    notes: Optional[str] = None  # Notlar
    conclusion: Optional[str] = None  # Sonuç (UYGUN/SAKINCALI)

# ===================== DYNAMIC FORM BUILDER API =====================

@app.get("/api/inspections/{inspection_id}/form")
async def get_inspection_form(inspection_id: str, current_user: User = Depends(get_current_user)):
    # Get inspection details
    query = {"id": inspection_id}
    if current_user.role == UserRole.DENETCI:
        query["inspector_id"] = current_user.id
    
    inspection = await db.inspections.find_one(query)
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")
    
    # Get customer info
    customer = await db.customers.find_one({"id": inspection["customer_id"]})
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    # Get equipment template
    equipment_type = inspection["equipment_info"].get("equipment_type", "CARASKAL")
    template = await db.equipment_templates.find_one({"equipment_type": equipment_type, "is_active": True})
    if not template:
        raise HTTPException(status_code=404, detail="Equipment template not found")
    
    # Get existing form data if any
    existing_data = inspection.get("report_data", {})
    
    # Prepare enhanced form data structure for Phase 6.2 & 6.3
    form_data = {
        "inspection_id": inspection_id,
        "equipment_type": equipment_type,
        "customer_name": customer["company_name"],
        "equipment_serial": inspection["equipment_info"].get("serial_number", ""),
        "control_items": template["categories"][0]["items"] if template["categories"] else [],
        "form_data": existing_data.get("form_data", {}),
        "general_info": existing_data.get("general_info", {}),
        "equipment_info": existing_data.get("equipment_info", {}),
        "photos": existing_data.get("photos", {}),
        "is_draft": existing_data.get("is_draft", True),
        "completion_percentage": existing_data.get("completion_percentage", 0),
        "last_saved": existing_data.get("last_saved"),
        "inspector": {
            "id": inspection["inspector_id"],
            "name": current_user.full_name if current_user.id == inspection["inspector_id"] else "Unknown"
        }
    }
    
    return form_data

@app.post("/api/inspections/{inspection_id}/form")
async def save_inspection_form(
    inspection_id: str,
    form_data: InspectionFormData,
    current_user: User = Depends(get_current_user)
):
    # Check authorization
    query = {"id": inspection_id}
    if current_user.role == UserRole.DENETCI:
        query["inspector_id"] = current_user.id
    
    inspection = await db.inspections.find_one(query)
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")
    
    # Validate form data
    if not form_data.form_results:
        raise HTTPException(status_code=400, detail="Form results cannot be empty")
    
    # Update inspection with form data
    update_data = {
        "report_data": form_data.dict(),
        "updated_at": datetime.utcnow()
    }
    
    await db.inspections.update_one({"id": inspection_id}, {"$set": update_data})
    
    # Get updated inspection
    updated_inspection = await db.inspections.find_one({"id": inspection_id})
    return {"message": "Form saved successfully", "inspection": Inspection(**updated_inspection)}

@app.put("/api/inspections/{inspection_id}/form")
async def update_inspection_form(
    inspection_id: str,
    form_data: dict,
    current_user: User = Depends(get_current_user)
):
    """Update inspection form data with enhanced Phase 6.2 & 6.3 features"""
    # Check authorization
    query = {"id": inspection_id}
    if current_user.role == UserRole.DENETCI:
        query["inspector_id"] = current_user.id
    
    inspection = await db.inspections.find_one(query)
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")
    
    # Prepare update data
    update_data = {
        "report_data": {
            "form_data": form_data.get("form_data", {}),
            "general_info": form_data.get("general_info", {}),
            "equipment_info": form_data.get("equipment_info", {}),
            "photos": form_data.get("photos", {}),
            "is_draft": form_data.get("is_draft", True),
            "completion_percentage": form_data.get("completion_percentage", 0),
            "last_saved": form_data.get("last_saved", datetime.utcnow().isoformat()),
            "saved_by": current_user.id
        },
        "updated_at": datetime.utcnow()
    }
    
    # If not draft, update status and completion time
    if not form_data.get("is_draft", True):
        update_data["status"] = "rapor_yazildi"
        update_data["completed_at"] = form_data.get("completed_at", datetime.utcnow().isoformat())
    
    await db.inspections.update_one({"id": inspection_id}, {"$set": update_data})
    
    # Get updated inspection
    updated_inspection = await db.inspections.find_one({"id": inspection_id})
    return {"message": "Form updated successfully", "inspection": Inspection(**updated_inspection)}

@app.get("/api/equipment-templates/{equipment_type}/form-structure")
async def get_equipment_form_structure(equipment_type: str, current_user: User = Depends(get_current_user)):
    """Get form structure for specific equipment type"""
    template = await db.equipment_templates.find_one({"equipment_type": equipment_type.upper(), "is_active": True})
    if not template:
        raise HTTPException(status_code=404, detail="Equipment template not found")
    
    # Format form structure for frontend
    form_structure = {
        "equipment_type": template["equipment_type"],
        "description": template["description"],
        "categories": []
    }
    
    for category in template["categories"]:
        category_data = {
            "code": category["code"],
            "name": category["name"],
            "items": [
                {
                    "id": item["id"],
                    "text": item["text"],
                    "input_type": item["input_type"],
                    "has_comment": item["has_comment"],
                    "required": item["required"],
                    "options": ["U", "UD", "U.Y"] if item["input_type"] == "dropdown" else None
                }
                for item in category["items"]
            ]
        }
        form_structure["categories"].append(category_data)
    
    return form_structure

@app.post("/api/equipment-templates/initialize")
async def initialize_caraskal_template(current_user: User = Depends(require_role(UserRole.ADMIN))):
    """Initialize Caraskal template in database"""
    # Check if Caraskal template already exists
    existing = await db.equipment_templates.find_one({"equipment_type": "CARASKAL"})
    if existing:
        return {"message": "Caraskal template already exists"}
    
    caraskal_data = get_caraskal_template()
    template_dict = caraskal_data.copy()
    template_dict["id"] = str(uuid.uuid4())
    template_dict["created_by"] = current_user.id
    template_dict["is_active"] = True
    template_dict["created_at"] = datetime.utcnow()
    template_dict["updated_at"] = datetime.utcnow()
    
    await db.equipment_templates.insert_one(template_dict)
    return {"message": "Caraskal template initialized successfully"}

# ===================== INSPECTION MANAGEMENT =====================

@app.post("/api/inspections", response_model=Inspection)
async def create_inspection(inspection: InspectionCreate, current_user: User = Depends(require_role(UserRole.PLANLAMA_UZMANI))):
    # Check for duplicate inspection (same customer + equipment combination)
    equipment_serial = inspection.equipment_info.get("serial_number")
    equipment_type = inspection.equipment_info.get("equipment_type")
    
    if equipment_serial and equipment_type:
        # Check for existing inspection with same customer and equipment
        existing_inspection = await db.inspections.find_one({
            "customer_id": inspection.customer_id,
            "equipment_info.serial_number": equipment_serial,
            "equipment_info.equipment_type": equipment_type,
            "status": {"$in": ["beklemede", "devam_ediyor", "rapor_yazildi"]}  # Only check active inspections
        })
        
        if existing_inspection:
            raise HTTPException(
                status_code=400, 
                detail=f"An active inspection already exists for this customer and equipment (Serial: {equipment_serial}, Type: {equipment_type})"
            )
    
    inspection_dict = inspection.dict()
    inspection_dict["id"] = str(uuid.uuid4())
    inspection_dict["created_by"] = current_user.id
    inspection_dict["created_at"] = datetime.utcnow()
    inspection_dict["updated_at"] = datetime.utcnow()
    inspection_dict["status"] = "beklemede"
    
    await db.inspections.insert_one(inspection_dict)
    return Inspection(**inspection_dict)

@app.get("/api/inspections", response_model=List[Inspection])
async def get_inspections(current_user: User = Depends(get_current_user)):
    query = {}
    if current_user.role == UserRole.DENETCI:
        query = {"inspector_id": current_user.id}
    
    inspections = await db.inspections.find(query).to_list(1000)
    return [Inspection(**inspection) for inspection in inspections]

@app.get("/api/inspections/pending-approval", response_model=List[Inspection])
async def get_pending_approval_inspections(current_user: User = Depends(require_role(UserRole.TEKNIK_YONETICI))):
    inspections = await db.inspections.find({"status": "rapor_yazildi"}).to_list(1000)
    return [Inspection(**inspection) for inspection in inspections]

@app.get("/api/inspections/{inspection_id}", response_model=Inspection)
async def get_inspection(inspection_id: str, current_user: User = Depends(get_current_user)):
    query = {"id": inspection_id}
    if current_user.role == UserRole.DENETCI:
        query["inspector_id"] = current_user.id
    
    inspection = await db.inspections.find_one(query)
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")
    return Inspection(**inspection)

# ===================== INSPECTION APPROVAL (TEKNİK YÖNETİCİ) =====================

@app.post("/api/inspections/{inspection_id}/approve")
async def approve_inspection(
    inspection_id: str, 
    approval: InspectionApproval,
    current_user: User = Depends(require_role(UserRole.TEKNIK_YONETICI))
):
    inspection = await db.inspections.find_one({"id": inspection_id})
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")
    
    if inspection["status"] != "rapor_yazildi":
        raise HTTPException(status_code=400, detail="Inspection is not ready for approval")
    
    update_data = {
        "approval_notes": approval.notes,
        "approved_by": current_user.id,
        "approved_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    if approval.action == "approve":
        update_data["status"] = "onaylandi"
    elif approval.action == "reject":
        update_data["status"] = "reddedildi"
    else:
        raise HTTPException(status_code=400, detail="Invalid action. Use 'approve' or 'reject'")
    
    await db.inspections.update_one({"id": inspection_id}, {"$set": update_data})
    
    updated_inspection = await db.inspections.find_one({"id": inspection_id})
    return Inspection(**updated_inspection)

# ===================== DASHBOARD STATS =====================

# ===================== INSPECTION UPDATE (DENETÇİ) =====================

@app.put("/api/inspections/{inspection_id}", response_model=Inspection)
async def update_inspection(
    inspection_id: str,
    inspection_update: InspectionUpdate,
    current_user: User = Depends(get_current_user)
):
    query = {"id": inspection_id}
    if current_user.role == UserRole.DENETCI:
        query["inspector_id"] = current_user.id
    
    inspection = await db.inspections.find_one(query)
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found or not authorized")
    
    update_data = {"updated_at": datetime.utcnow()}
    if inspection_update.status:
        update_data["status"] = inspection_update.status
    if inspection_update.report_data:
        update_data["report_data"] = inspection_update.report_data
    
    await db.inspections.update_one({"id": inspection_id}, {"$set": update_data})
    
    updated_inspection = await db.inspections.find_one({"id": inspection_id})
    return Inspection(**updated_inspection)

@app.get("/api/dashboard/stats")
async def get_dashboard_stats(current_user: User = Depends(get_current_user)):
    stats = {}
    
    if current_user.role in [UserRole.ADMIN, UserRole.PLANLAMA_UZMANI]:
        total_customers = await db.customers.count_documents({})
        total_inspections = await db.inspections.count_documents({})
        pending_inspections = await db.inspections.count_documents({"status": "beklemede"})
        completed_inspections = await db.inspections.count_documents({"status": "onaylandi"})
        
        stats.update({
            "total_customers": total_customers,
            "total_inspections": total_inspections,
            "pending_inspections": pending_inspections,
            "completed_inspections": completed_inspections
        })
    
    if current_user.role == UserRole.DENETCI:
        my_inspections = await db.inspections.count_documents({"inspector_id": current_user.id})
        my_pending = await db.inspections.count_documents({"inspector_id": current_user.id, "status": "beklemede"})
        my_in_progress = await db.inspections.count_documents({"inspector_id": current_user.id, "status": "devam_ediyor"})
        my_completed = await db.inspections.count_documents({"inspector_id": current_user.id, "status": "rapor_yazildi"})
        
        stats.update({
            "my_inspections": my_inspections,
            "my_pending": my_pending,
            "my_in_progress": my_in_progress,
            "my_completed": my_completed
        })
    
    if current_user.role == UserRole.TEKNIK_YONETICI:
        pending_approval = await db.inspections.count_documents({"status": "rapor_yazildi"})
        approved_today = await db.inspections.count_documents({
            "status": "onaylandi",
            "approved_by": current_user.id,
            "approved_at": {"$gte": datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)}
        })
        rejected_reports = await db.inspections.count_documents({"status": "reddedildi"})
        total_approved = await db.inspections.count_documents({"approved_by": current_user.id, "status": "onaylandi"})
        
        stats.update({
            "pending_approval": pending_approval,
            "approved_today": approved_today,
            "rejected_reports": rejected_reports,
            "total_approved": total_approved
        })
    
    return stats

# ===================== STARTUP EVENT =====================

@app.on_event("startup")
async def startup_event():
    # Create default admin user if doesn't exist
    admin_exists = await db.users.find_one({"username": "admin"})
    if not admin_exists:
        admin_user = {
            "id": str(uuid.uuid4()),
            "username": "admin",
            "email": "admin@royalcert.com",
            "full_name": "System Administrator",
            "role": UserRole.ADMIN,
            "password": get_password_hash("admin123"),
            "is_active": True,
            "created_at": datetime.utcnow()
        }
        await db.users.insert_one(admin_user)
        print("✅ Default admin user created: admin/admin123")
    
    # Initialize Caraskal template if doesn't exist
    caraskal_exists = await db.equipment_templates.find_one({"equipment_type": "CARASKAL"})
    if not caraskal_exists:
        admin_user = await db.users.find_one({"role": UserRole.ADMIN})
        if admin_user:
            caraskal_data = get_caraskal_template()
            template_dict = caraskal_data.copy()
            template_dict["id"] = str(uuid.uuid4())
            template_dict["created_by"] = admin_user["id"]
            template_dict["is_active"] = True
            template_dict["created_at"] = datetime.utcnow()
            template_dict["updated_at"] = datetime.utcnow()
            
            await db.equipment_templates.insert_one(template_dict)
            print("✅ Caraskal template initialized automatically")

# ===================== TEMPLATE UPLOAD & WORD PARSING =====================

def parse_word_document(file_content: bytes, filename: str) -> dict:
    """Parse Word document to extract inspection template structure"""
    try:
        # Read Word document
        doc = Document(io.BytesIO(file_content))
        
        # Extract document text
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        
        # Extract table data for better structure parsing
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        
        # Determine equipment type and template type from filename
        equipment_type = "UNKNOWN"
        template_type = "FORM"  # Default to form
        
        filename_upper = filename.upper()
        
        # Extract equipment type (ignore MUAYENE, FORMU, RAPORU keywords)
        if "FORKL" in filename_upper:
            equipment_type = "FORKLIFT"
        elif "CARASKAL" in filename_upper:
            equipment_type = "CARASKAL"
        elif "ISKELE" in filename_upper:
            equipment_type = "ISKELE"
        elif "VINC" in filename_upper or "CRANE" in filename_upper:
            equipment_type = "VINC"
        elif "ASANSÖR" in filename_upper or "ELEVATOR" in filename_upper:
            equipment_type = "ASANSÖR"
        
        # Determine template type
        if "RAPOR" in filename_upper:
            template_type = "REPORT"
        elif "FORM" in filename_upper:
            template_type = "FORM"
        
        # SMART Parse control items from text - GET REAL CONTROL ITEMS ONLY
        control_items = []
        categories = {}
        
        # Look for numbered items (1., 2., 3., etc.) in text - BETTER PATTERN
        item_pattern = r'^(\d+)\.\s*(.+)$'
        matches = re.findall(item_pattern, text, re.MULTILINE)
        
        # Smart filtering to get REAL control items only
        valid_matches = []
        seen_texts = set()  # Prevent duplicates
        
        for match in matches:
            item_number = int(match[0])
            item_text = match[1].strip()
            
            # Skip if item number is unreasonable (but not too restrictive)
            if item_number < 1 or item_number > 100:  # Allow up to 100 for complex equipment
                continue
                
            # Skip if item text is too short (real control items are descriptive)
            if len(item_text) < 15:  # Real control items are at least 15 chars
                continue
                
            # Skip common headers and non-control text (SMART FILTERING)
            skip_patterns = [
                'GENEL', 'BİLGİLER', 'MUAYENE', 'KONTROL', 'ETİKET', 'TEST', 'FORM', 'RAPOR', 
                'BAŞLIK', 'TABLE', 'DEĞER', 'DURUM', 'TARİH', 'NO', 'ADI', 'KODU', 'MARKASı',
                'TİPİ', 'SERİ', 'İMAL', 'YIL', 'KAPASITE', 'YÜKSEKLIK', 'MESAFE', 'ÖLÇÜM',
                'DEĞERLENDİRME', 'AÇIKLAMA', 'NOT'
            ]
            
            # Skip if item text is primarily a header/label
            if any(pattern in item_text.upper() for pattern in skip_patterns):
                continue
                
            # Skip if text contains primarily numbers/codes/values (not control descriptions)  
            if re.search(r'^\d+[.\-/\s]*\d*$', item_text.strip()):
                continue
                
            # Skip repetitive or too similar texts
            text_key = re.sub(r'\s+', ' ', item_text.lower())
            if text_key in seen_texts:
                continue
            seen_texts.add(text_key)
            
            # Skip if text is primarily symbols or formatting
            if len(re.sub(r'[^a-zA-ZğüşıöçĞÜŞİÖÇ]', '', item_text)) < 10:
                continue
            
            # This looks like a REAL control item!
            valid_matches.append((item_number, item_text))
        
        # Sort by item number (preserve original order)
        valid_matches.sort(key=lambda x: x[0])
        
        print(f"DEBUG: Found {len(valid_matches)} valid control items after smart filtering")
        
        current_category = 'A'
        
        for match in valid_matches:
            item_number = int(match[0])
            item_text = match[1].strip()
            
            # Smart category determination based on typical equipment inspection structure
            if item_number <= 12:
                current_category = 'A'  # Usually control/cabin systems
            elif item_number <= 20:
                current_category = 'B'  # Usually movement/drive systems  
            elif item_number <= 27:
                current_category = 'C'  # Usually indicators/warnings
            elif item_number <= 35:
                current_category = 'D'  # Usually braking systems
            elif item_number <= 42:
                current_category = 'E'  # Usually lifting/chains
            elif item_number <= 48:
                current_category = 'F'  # Usually forks/attachments
            elif item_number <= 55:
                current_category = 'G'  # Usually fuel/emissions
            else:
                current_category = 'H'  # Usually other controls/safety
            
            # Add to categories count
            if current_category not in categories:
                categories[current_category] = 0
            categories[current_category] += 1
            
            control_items.append({
                "id": item_number,
                "text": item_text,
                "category": current_category,
                "has_dropdown": True,  # All items have U/UD/U.Y dropdown
                "has_comment": True,   # All items have comment field
                "has_photo": True      # All items can have photos
            })
        
        # FALLBACK: If no numbered items found, try smarter table parsing
        if not control_items and tables:
            print("DEBUG: No numbered items found, trying smart table parsing...")
            item_id = 1
            seen_texts = set()
            
            for table in tables:
                for row in table:
                    for cell_text in row:
                        # Smart table cell filtering for REAL control items
                        if (len(cell_text) > 20 and len(cell_text) < 300 and  # Reasonable length
                            not any(x in cell_text.upper() for x in ['GENEL', 'BİLGİLER', 'MUAYENE', 'TEST', 'ETİKET', 
                                                                     'KONTROL', 'FORM', 'RAPOR', 'TABLE', 'BAŞLIK',
                                                                     'NO', 'ADI', 'KODU', 'DURUM', 'TARİH']) and
                            not cell_text.upper().strip() in ['D', 'U', 'UD', 'U.Y'] and
                            cell_text.count('.') < 5 and  # Avoid dotted number sequences
                            len(re.sub(r'[^a-zA-ZğüşıöçĞÜŞİÖÇ]', '', cell_text)) > 15):  # Has enough letters
                            
                            # Check for duplicates
                            text_key = re.sub(r'\s+', ' ', cell_text.lower().strip())
                            if text_key not in seen_texts:
                                seen_texts.add(text_key)
                                
                                # Smart category assignment
                                if item_id <= 12:
                                    current_category = 'A'
                                elif item_id <= 20:
                                    current_category = 'B'
                                elif item_id <= 27:
                                    current_category = 'C'
                                elif item_id <= 35:
                                    current_category = 'D'
                                elif item_id <= 42:
                                    current_category = 'E'
                                elif item_id <= 48:
                                    current_category = 'F'
                                elif item_id <= 55:
                                    current_category = 'G'
                                else:
                                    current_category = 'H'
                                
                                control_items.append({
                                    "id": item_id,
                                    "text": cell_text.strip(),
                                    "category": current_category,
                                    "has_dropdown": True,
                                    "has_comment": True,
                                    "has_photo": True
                                })
                                item_id += 1
                                
                                # Reasonable limit - but not too restrictive
                                if item_id > 80:  # Allow up to 80 for complex equipment
                                    break
            
            print(f"DEBUG: Table parsing found {len(control_items)} control items")
        
        # Group control items by category
        categories_dict = {}
        for item in control_items:
            category = item.get('category', 'A')
            if category not in categories_dict:
                categories_dict[category] = []
            categories_dict[category].append({
                "id": item["id"],
                "text": item["text"],
                "category": category,
                "input_type": "dropdown",
                "has_comment": True,
                "required": True
            })
        
        # Create category structure
        categories_list = []
        category_names = {
            'A': 'KATEGORI A',
            'B': 'KATEGORI B', 
            'C': 'KATEGORI C',
            'D': 'KATEGORI D',
            'E': 'KATEGORI E',
            'F': 'KATEGORI F',
            'G': 'KATEGORI G',
            'H': 'KATEGORI H'
        }
        
        for category_code in sorted(categories_dict.keys()):
            categories_list.append({
                "code": category_code,
                "name": category_names.get(category_code, f"KATEGORI {category_code}"),
                "items": categories_dict[category_code]
            })
        
        # Create template structure
        # Generate template name based on equipment type and template type
        if template_type == "REPORT":
            template_name = f"{equipment_type} MUAYENE RAPORU"
            description = f"{equipment_type} ekipmanı için otomatik PDF rapor template'i"
        else:
            template_name = f"{equipment_type} MUAYENE FORMU"
            description = f"{equipment_type} ekipmanı için denetim kontrol formu"
        
        template_data = {
            "name": template_name,
            "equipment_type": equipment_type,
            "template_type": template_type,  # FORM or REPORT
            "description": description,
            "categories": categories_list,
            "total_items": len(control_items),
            "parsed_from": filename,
            "parse_date": datetime.utcnow(),
            "is_active": True
        }
        
        return template_data
        
    except Exception as e:
        print(f"Error parsing Word document {filename}: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to parse Word document: {str(e)}")

@app.post("/api/equipment-templates/upload")
async def upload_template_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Upload and parse Word document to create equipment template"""
    
    # Check authorization (only admin can upload templates)
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Only admins can upload templates")
    
    # Validate file type
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only Word documents (.docx, .doc) are supported")
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Parse Word document
        template_data = parse_word_document(file_content, file.filename)
        
        # Add metadata
        template_data["id"] = str(uuid.uuid4())
        template_data["created_by"] = current_user.id
        template_data["created_at"] = datetime.utcnow()
        template_data["updated_at"] = datetime.utcnow()
        
        # Check if template with same equipment_type and template_type already exists (only active templates)
        existing_template = await db.equipment_templates.find_one({
            "equipment_type": template_data["equipment_type"],
            "template_type": template_data["template_type"],
            "is_active": True
        })
        
        if existing_template:
            raise HTTPException(
                status_code=400, 
                detail=f"{template_data['equipment_type']} {template_data['template_type']} template already exists"
            )
        
        # Insert template into database
        await db.equipment_templates.insert_one(template_data)
        
        return {
            "message": "Template uploaded and parsed successfully",
            "template": {
                "id": template_data["id"],
                "name": template_data["name"],
                "equipment_type": template_data["equipment_type"],
                "template_type": template_data["template_type"],
                "total_items": template_data["total_items"],
                "categories": len(template_data["categories"])
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Template upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Template upload failed: {str(e)}")

@app.post("/api/equipment-templates/bulk-upload")
async def bulk_upload_templates(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user)
):
    """Bulk upload multiple Word documents to create equipment templates"""
    
    # Check authorization
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Only admins can upload templates")
    
    if len(files) > 50:
        raise HTTPException(status_code=400, detail="Maximum 50 files can be uploaded at once")
    
    results = {
        "successful": [],
        "failed": [],
        "total_files": len(files)
    }
    
    for file in files:
        try:
            # Validate file type
            if not file.filename.endswith(('.docx', '.doc')):
                results["failed"].append({
                    "filename": file.filename,
                    "error": "Only Word documents (.docx, .doc) are supported"
                })
                continue
            
            # Read and parse file
            file_content = await file.read()
            template_data = parse_word_document(file_content, file.filename)
            
            # Add metadata
            template_data["id"] = str(uuid.uuid4())
            template_data["created_by"] = current_user.id
            template_data["created_at"] = datetime.utcnow()
            template_data["updated_at"] = datetime.utcnow()
            
            # Check if template already exists (only active templates)
            existing_template = await db.equipment_templates.find_one({
                "equipment_type": template_data["equipment_type"],
                "template_type": template_data["template_type"],
                "is_active": True
            })
            
            if existing_template:
                results["failed"].append({
                    "filename": file.filename,
                    "error": f"{template_data['equipment_type']} {template_data['template_type']} template already exists"
                })
                continue
            
            # Insert template
            await db.equipment_templates.insert_one(template_data)
            
            results["successful"].append({
                "filename": file.filename,
                "template_name": template_data["name"],
                "equipment_type": template_data["equipment_type"],
                "template_type": template_data["template_type"],
                "total_items": template_data["total_items"]
            })
            
        except Exception as e:
            results["failed"].append({
                "filename": file.filename,
                "error": str(e)
            })
    
    return {
        "message": f"Bulk upload completed. {len(results['successful'])} successful, {len(results['failed'])} failed",
        "results": results
    }

# ===================== DATA REPAIR UTILITIES =====================

@app.post("/api/fix/orphaned-inspector-ids")
async def fix_orphaned_inspector_ids(current_user: User = Depends(get_current_user)):
    """Fix orphaned inspector IDs in inspections"""
    
    # Only admin can run data fixes
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Only admins can run data fixes")
    
    try:
        # Get all active inspectors
        inspectors = await db.users.find({"role": "denetci"}).to_list(None)
        # Filter for active inspectors (handle missing is_active field)
        inspectors = [insp for insp in inspectors if insp.get("is_active", True)]
        print(f"DEBUG: Found {len(inspectors)} inspectors in data fix endpoint")
        valid_inspector_ids = {insp["id"] for insp in inspectors}
        
        # Get all beklemede inspections
        beklemede_inspections = await db.inspections.find({"status": "beklemede"}).to_list(None)
        
        fixed_count = 0
        results = []
        
        for inspection in beklemede_inspections:
            inspector_id = inspection.get("inspector_id")
            inspection_id = inspection.get("id")
            equipment_type = inspection.get("equipment_info", {}).get("equipment_type", "N/A")
            
            if inspector_id not in valid_inspector_ids:
                # Orphaned inspector_id found - reassign to first available inspector
                if inspectors:
                    new_inspector = inspectors[0]  # Assign to first available inspector
                    
                    # Update the inspection
                    await db.inspections.update_one(
                        {"id": inspection_id},
                        {
                            "$set": {
                                "inspector_id": new_inspector["id"],
                                "updated_at": datetime.utcnow()
                            }
                        }
                    )
                    
                    results.append({
                        "inspection_id": inspection_id,
                        "equipment_type": equipment_type,
                        "old_inspector_id": inspector_id,
                        "new_inspector_id": new_inspector["id"],
                        "new_inspector_name": new_inspector["full_name"],
                        "action": "reassigned"
                    })
                    fixed_count += 1
                else:
                    results.append({
                        "inspection_id": inspection_id,
                        "equipment_type": equipment_type,
                        "old_inspector_id": inspector_id,
                        "action": "no_inspectors_available"
                    })
            else:
                results.append({
                    "inspection_id": inspection_id,
                    "equipment_type": equipment_type,
                    "inspector_id": inspector_id,
                    "action": "valid_assignment"
                })
        
        return {
            "message": f"Data repair completed. Fixed {fixed_count} orphaned inspector IDs",
            "total_inspections_checked": len(beklemede_inspections),
            "total_inspectors_available": len(inspectors),
            "fixed_count": fixed_count,
            "results": results
        }
        
    except Exception as e:
        print(f"Data fix error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Data fix failed: {str(e)}")

# ===================== HEALTH CHECK =====================

@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow()}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)